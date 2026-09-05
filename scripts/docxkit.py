# -*- coding: utf-8 -*-
"""Общие помощники для сборки договорных документов в DOCX."""
import os

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "docs", "contract")

BLANK = "_" * 30
BLANK_S = "_" * 14
DOC_REF = "к Договору № " + BLANK_S + " от «___» __________ 20___ г."


def new_doc():
    d = Document()
    for s in d.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(1.5)
    st = d.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(11)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = st.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _page_numbers(d)
    return d


def _page_numbers(d):
    p = d.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, end):
        run._r.append(el)
    run.font.size = Pt(9)
    run.font.name = "Times New Roman"


def title(d, text, sub=None):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    if sub:
        p2 = d.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(12)
        p2.add_run(sub).font.size = Pt(11)


def h(d, text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    p.add_run(text).bold = True
    return p


def para(d, text, indent=0, bold=False, italic=False,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=None):
    p = d.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    return p


def bullet(d, text):
    p = d.add_paragraph(text, style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)
    return p


def section(d, num, heading, items):
    """Раздел с автонумерацией пунктов num.1, num.2 ..."""
    h(d, "%d. %s" % (num, heading))
    n = 0
    for item in items:
        if isinstance(item, tuple):
            kind, text = item
            if kind == "sub":
                para(d, text, indent=1.0)
            elif kind == "note":
                para(d, text, indent=1.0, italic=True, size=10)
            continue
        n += 1
        para(d, "%d.%d. %s" % (num, n, item))


def table(d, headers, rows, widths=None, font=9.5):
    t = d.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htxt in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.paragraphs[0].text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(htxt)
        r.bold = True
        r.font.size = Pt(font)
        r.font.name = "Times New Roman"
    trPr = t.rows[0]._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:tblHeader"))
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].paragraphs[0].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(val))
            r.font.size = Pt(font)
            r.font.name = "Times New Roman"
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    d.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def no_borders(tbl):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        borders.append(el)
    tbl._tbl.tblPr.append(borders)


def place_and_date(d):
    tbl = d.add_table(rows=1, cols=2)
    left, right = tbl.rows[0].cells
    left.paragraphs[0].add_run("г. " + BLANK_S).font.size = Pt(11)
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.add_run("«___» __________ 20___ г.").font.size = Pt(11)
    for c in (left, right):
        for p in c.paragraphs:
            for r in p.runs:
                r.font.name = "Times New Roman"
    no_borders(tbl)
    d.add_paragraph()


# --- Реквизиты Исполнителя (заполнены) ------------------------------------
EXECUTOR_NAME = "Яковлев Артём Сергеевич"
EXECUTOR_SHORT = "Яковлев А. С."
EXECUTOR_INN = "532123260012"
EXECUTOR_ADDR = ("Новгородская область, г. Великий Новгород, "
                 "ул. 8 Марта, д. 27, кв. 36")
EXECUTOR_ACC = "40817810943861695557"
EXECUTOR_BANK = "Архангельское отделение № 8637 ПАО Сбербанк"
EXECUTOR_BIK = "041117601"
EXECUTOR_CORR = "30101810100000000601"
EXECUTOR_PHONE = "+7 911 609-63-41"
EXECUTOR_EMAIL = "jackowleffa@yandex.ru"
EXECUTOR_BASIS = ("справки о постановке на учёт физического лица в качестве "
                  "налогоплательщика налога на профессиональный доход")
NPD_NOTE = ("НДС не облагается в связи с применением Исполнителем специального "
            "налогового режима «Налог на профессиональный доход» "
            "(Федеральный закон от 27.11.2018 № 422-ФЗ)")


def parties(d, customer_role="Заказчик", verb="заключили настоящий Договор о нижеследующем."):
    para(
        d,
        "{c}, именуем___ в дальнейшем «{cr}», в лице {p1}, действующего на основании {b1}, "
        "с одной стороны, и гражданин Российской Федерации {name}, ИНН {inn}, применяющий "
        "специальный налоговый режим «Налог на профессиональный доход», именуемый "
        "в дальнейшем «Исполнитель», действующий от своего имени на основании {basis}, "
        "с другой стороны, совместно именуемые «Стороны», а по отдельности — «Сторона», "
        "{verb}".format(
            c=BLANK, cr=customer_role, p1=BLANK_S, b1=BLANK_S,
            name=EXECUTOR_NAME, inn=EXECUTOR_INN, basis=EXECUTOR_BASIS, verb=verb,
        ),
    )


REQUISITES = [
    "Наименование: " + BLANK,
    "ИНН: " + BLANK_S + "   КПП: " + BLANK_S,
    "ОГРН (ОГРНИП): " + BLANK,
    "Юридический адрес: " + BLANK,
    "Почтовый адрес: " + BLANK,
    "Расчётный счёт: " + BLANK,
    "Банк: " + BLANK,
    "БИК: " + BLANK_S + "  К/с: " + BLANK_S,
    "Телефон: " + BLANK_S,
    "E-mail: " + BLANK,
    "",
    "Должность: " + BLANK,
    "",
    "_____________ / " + BLANK_S + " /",
    "                                      м.п.",
]

EXECUTOR_REQUISITES = [
    EXECUTOR_NAME,
    "Налогоплательщик налога на профессиональный доход (самозанятый)",
    "ИНН: " + EXECUTOR_INN,
    "Адрес регистрации: " + EXECUTOR_ADDR,
    "Почтовый адрес: тот же",
    "Счёт: " + EXECUTOR_ACC,
    "Банк: " + EXECUTOR_BANK,
    "БИК: " + EXECUTOR_BIK + "  К/с: " + EXECUTOR_CORR,
    "Телефон: " + EXECUTOR_PHONE,
    "E-mail: " + EXECUTOR_EMAIL,
    "",
    "",
    "",
    "_____________ / " + EXECUTOR_SHORT + " /",
    "",
]

SIGN_ONLY = [
    "Должность: " + BLANK,
    "",
    "",
    "_____________ / " + BLANK_S + " /",
    "                                      м.п.",
]

EXECUTOR_SIGN_ONLY = [
    EXECUTOR_NAME,
    "",
    "",
    "_____________ / " + EXECUTOR_SHORT + " /",
    "",
]


def signatures(d, full=True, heading=None):
    h(d, heading or ("РЕКВИЗИТЫ И ПОДПИСИ СТОРОН" if full else "ПОДПИСИ СТОРОН"))
    left_body = REQUISITES if full else SIGN_ONLY
    right_body = EXECUTOR_REQUISITES if full else EXECUTOR_SIGN_ONLY
    tbl = d.add_table(rows=1, cols=2)
    for i, cell in enumerate(tbl.rows[0].cells):
        cell.width = Cm(8.0)
        cell.paragraphs[0].text = ""
        r = cell.paragraphs[0].add_run("ЗАКАЗЧИК" if i == 0 else "ИСПОЛНИТЕЛЬ")
        r.bold = True
        r.font.size = Pt(10)
        r.font.name = "Times New Roman"
        for line in (left_body if i == 0 else right_body):
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            rr = p.add_run(line)
            rr.font.size = Pt(10)
            rr.font.name = "Times New Roman"
    no_borders(tbl)


def appendix(num, name, ref=DOC_REF):
    d = new_doc()
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run("Приложение № %s\n" % num)
    r2 = p.add_run(ref)
    for r in (r1, r2):
        r.font.size = Pt(10)
        r.italic = True
    d.add_paragraph()
    title(d, name)
    return d


def save(d, filename):
    os.makedirs(OUT, exist_ok=True)
    path = os.path.join(OUT, filename)
    d.save(path)
    print("OK:", filename)
    return path
